from __future__ import annotations

import asyncio
import math

from markitdown import MarkItDown

from src.services.document_database.ocr.doc_router import ProcessDecision
from src.services.document_database.ocr.page_format import format_document_with_pages
from src.services.document_database.ocr.types.interface import FileReader


class DocxReader(FileReader):
    async def read(
        self,
        filename: str,
        path: str,
        fate: ProcessDecision,
    ) -> str:
        if fate == ProcessDecision.OCR:
            # DOCX files don't require OCR
            raise ValueError("OCR is not applicable for DOCX files")

        file_path = self._resolve_file_path(path, filename)

        try:
            page_texts = await asyncio.to_thread(
                self._extract_docx_pages_sync, str(file_path)
            )
            if len(page_texts) <= 1:
                estimated_pages = self.get_page_count(
                    filename=filename, path=str(file_path)
                )
                if estimated_pages > 1:
                    combined = page_texts[0] if page_texts else ""
                    page_texts = self._split_text_into_estimated_pages(
                        combined, estimated_pages
                    )
            return format_document_with_pages(filename, page_texts)
        except Exception:
            md = MarkItDown(enable_plugins=False)  # Set to True to enable plugins
            result = await asyncio.to_thread(md.convert, str(file_path))
            fallback_text = (
                result.text_content if result and result.text_content else ""
            )
            return format_document_with_pages(filename, [fallback_text])

    def _extract_docx_pages_sync(self, file_path: str) -> list[str]:
        from docx import Document

        document = Document(file_path)
        pages: list[list[str]] = [[]]

        for paragraph in document.paragraphs:
            paragraph_text = (paragraph.text or "").strip()
            if paragraph_text:
                pages[-1].append(paragraph_text)

            page_breaks = self._count_page_breaks(paragraph)
            for _ in range(page_breaks):
                pages.append([])

        collapsed = ["\n".join(page).strip() for page in pages]
        while len(collapsed) > 1 and not collapsed[-1]:
            collapsed.pop()
        return collapsed or [""]

    def _count_page_breaks(self, paragraph) -> int:
        page_breaks = 0
        for run in paragraph.runs:
            run_element = run._element
            page_breaks += len(
                run_element.xpath(
                    ".//*[local-name()='br' and @*[local-name()='type']='page']"
                )
            )
            page_breaks += len(
                run_element.xpath(".//*[local-name()='lastRenderedPageBreak']")
            )
        return page_breaks

    def _split_text_into_estimated_pages(
        self, full_text: str, estimated_pages: int
    ) -> list[str]:
        if estimated_pages <= 1:
            return [full_text]

        words = full_text.split()
        if not words:
            return [""]

        chunk_size = max(1, math.ceil(len(words) / estimated_pages))
        chunks = [
            " ".join(words[i : i + chunk_size]).strip()
            for i in range(0, len(words), chunk_size)
        ]
        return chunks or [full_text]

    def get_page_count(self, filename: str, path: str) -> int:
        """
        Estimate page count for DOCX files.

        Uses a heuristic approach:
        1. Try python-docx to count words (~500 words/page)
        2. Fall back to MarkItDown conversion and estimate from character count
        3. Default to 1 if all methods fail

        Note: Accurate page count requires rendering, so this is an approximation.
        """
        file_path = self._resolve_file_path(path, filename)
        try:
            from docx import Document

            doc = Document(str(file_path))

            # Estimate: ~500 words per page, ~5 chars per word
            total_chars = sum(len(paragraph.text) for paragraph in doc.paragraphs)
            estimated_words = total_chars / 5
            estimated_pages = max(1, int(estimated_words / 500))

            return estimated_pages
        except ImportError:
            # If python-docx not installed, fall back to MarkItDown conversion
            try:
                md = MarkItDown(enable_plugins=False)
                result = md.convert(str(file_path))
                # Rough estimate: 3000 chars per page
                estimated_pages = max(1, len(result.text_content) // 3000)
                return estimated_pages
            except Exception:
                return 1
        except Exception:
            return 1
